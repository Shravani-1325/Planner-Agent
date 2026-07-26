import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(title: str, sections: list, output_dir: str = "output") -> str:
    os.makedirs(output_dir, exist_ok=True)
    doc = Document() # creates blank in memory word doc objest

    heading = doc.add_heading(title, level = 0) # Level 0 - biggest font, level 1 - normal, level 2 - subsection
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    
    doc.add_paragraph() # spacer
    
    for section in sections:
        doc.add_heading(section.get("heading", "Section"), level = 1)
        doc.add_paragraph(section.get("content", ""))

    safe_name = "".join(c if c.isalnum() or c in (" ", "_") else "" for c in title)
    safe_name = safe_name.strip().replace(" ", "_") or "document"

    filename = f"{safe_name}_{datetime.now().strftime("%Y%m%d_H%M%S")}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath

        
        
    
    
    
    
    